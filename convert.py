#!/usr/bin/env python3

import argparse

from typing import Optional, Union

from docx import Document


class Element:
    pass


class Paragraph(Element):
    text: str
    number: Optional[str]

    def __init__(self, text, number):
        self.text = text
        self.number = number

    def __repr__(self):
        return f"<Paragraph: {self.text}>"


class List(Element):
    items = list[Union[Element, str]]


class Article:
    number: str
    heading: str
    paragraphs: list[Paragraph]

    def __init__(self, number, heading, paragraphs):
        self.number = number
        self.heading = heading
        self.paragraphs = paragraphs

    def __repr__(self):
        return f"<Article: {self.heading} ({self.number})>"


class Chapter:
    number: str
    heading: str
    articles: list[Article]

    def __init__(self, number, heading):
        self.number = number
        self.heading = heading
        self.articles = []

    def __repr__(self):
        return f"<Chapter: {self.heading} ({self.number})>"


def get_chapter_number(text):
    return text.upper().replace("CHAPTER ", "")


def assert_is_blank_line(paragraphs):
    text = next(paragraphs).text

    assert not text, f"{text} is not a blank line"


def get_chapter(paragraphs):
    chapter_number_text = next(paragraphs).text
    assert chapter_number_text.upper().startswith("CHAPTER"), f"{chapter_number_text} does not start with CHAPTER"
    chapter_number = get_chapter_number(chapter_number_text)

    assert_is_blank_line(paragraphs)

    chapter_heading = next(paragraphs).text

    chapter = Chapter(chapter_number, chapter_heading)

    return chapter


def get_article_number(text):
    return text.replace("Article ", "").strip()


def is_blank_line(text):
    return not text.strip()


def get_articles(paragraphs):
    articles = []

    while True:
        try:
            article_number_text = next(paragraphs).text
        except StopIteration:
            break

        assert article_number_text.startswith("Article"), f"{article_number_text} does not start with 'Article'"
        article_number = get_article_number(article_number_text)

        article_heading = next(paragraphs).text.strip()

        assert is_blank_line(next(paragraphs).text)

        last_line = ""

        article_paragraphs = []
        while True:
            try:
                line = next(paragraphs).text.strip()
            except StopIteration:
                break

            if not last_line and not line:
                break

            if line:
                paragraph = Paragraph(line, None)
                article_paragraphs.append(paragraph)

            last_line = line

        article = Article(article_number, article_heading, article_paragraphs)
        print(article)
        for paragraph in article.paragraphs:
            print(paragraph)
        articles.append(article)

    return articles


def convert(file):
    document = Document(file)
    paragraphs = iter(document.paragraphs)

    chapter = get_chapter(paragraphs)
    print(chapter)

    # Consume a blank line
    assert is_blank_line(next(paragraphs).text)
    # Consume a blank line
    assert is_blank_line(next(paragraphs).text)

    articles = get_articles(paragraphs)
    chapter.articles = articles


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert Word doc to markdown")
    parser.add_argument("document", type=argparse.FileType("rb"))
    args = parser.parse_args()

    convert(args.document)
